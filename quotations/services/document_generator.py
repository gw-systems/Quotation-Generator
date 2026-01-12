"""
Document generation service for creating DOCX quotations
"""
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.conf import settings
from decimal import Decimal


class QuotationDocxGenerator:
    """Generate DOCX quotations from template"""
    
    def __init__(self, quotation):
        """
        Initialize generator with quotation instance
        
        Args:
            quotation: Quotation model instance
        """
        self.quotation = quotation
        self.client = quotation.client
        self.template_path = settings.QUOTATION_TEMPLATE_PATH
        
    def generate(self, output_path=None):
        """
        Generate DOCX document
        
        Args:
            output_path: Optional custom output path
            
        Returns:
            str: Path to generated DOCX file
        """
        # Load template
        doc = Document(self.template_path)
        
        # Populate document
        self._populate_client_details(doc)
        self._populate_quotation_summary(doc)
        self._populate_pricing_sections(doc)
        
        # Determine output path
        if not output_path:
            output_dir = os.path.join(settings.MEDIA_ROOT, 'quotations', 'docx')
            os.makedirs(output_dir, exist_ok=True)
            filename = f"{self.quotation.quotation_number}.docx"
            output_path = os.path.join(output_dir, filename)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _populate_client_details(self, doc):
        """Populate client details in document tables"""
        tables = doc.tables
        if len(tables) > 0:
            client_table = tables[0]
            
            try:
                client_table.rows[0].cells[0].text = f"Client Name: {self.client.client_name}"
                client_table.rows[0].cells[1].text = f"Company Name: {self.client.company_name}"
                
                client_table.rows[1].cells[0].text = f"Email: {self.client.email}"
                client_table.rows[1].cells[1].text = f"Contact Number: {self.client.contact_number}"
                
                client_table.rows[2].cells[0].text = f"Address: {self.client.address}"
                if len(client_table.rows[2].cells) > 1:
                    client_table.rows[2].cells[1].text = ""
            except IndexError:
                pass
    
    def _populate_quotation_summary(self, doc):
        """Populate quotation summary in document tables"""
        tables = doc.tables
        if len(tables) > 1:
            summary_table = tables[1]
            
            try:
                # Row 0: Date (Cell 0)
                summary_table.rows[0].cells[0].text = f"Date: {self.quotation.date.strftime('%d-%m-%Y')}"
                
                # Row 1: Validity (Cell 0) | Point of Contact (Cell 1)
                # Check row 1 lengths to be safe
                if len(summary_table.rows) > 1:
                    row1 = summary_table.rows[1]
                    if len(row1.cells) > 0:
                        row1.cells[0].text = f"Validity Period: {self.quotation.validity_period} days (Until {self.quotation.validity_date.strftime('%d-%m-%Y')})"
                    
                    if len(row1.cells) > 1:
                        row1.cells[1].text = f"Point of Contact: {self.quotation.point_of_contact}"
            except IndexError:
                pass
    
    def _find_pricing_header_index(self, doc):
        """Find the paragraph index of 'PRICING DETAILS' header"""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.upper()
            if 'PRICING' in text and ('DETAILS' in text or 'NCR' in text or 'REGION' in text):
                return i
        return -1
    
    def _find_pricing_table_index(self, tables):
        """Find the index of the pricing table (4 columns, 10+ rows)"""
        for i, table in enumerate(tables):
            if len(table.columns) == 4 and len(table.rows) >= 8:
                return i
        # Fallback to table 2
        if len(tables) > 2:
            return 2
        return -1
    
    def _populate_pricing_sections(self, doc):
        """Populate pricing sections for all locations"""
        locations = list(self.quotation.locations.all().order_by('order'))
        if not locations:
            return
        
        tables = doc.tables
        
        # Find the pricing header paragraph
        pricing_header_idx = self._find_pricing_header_index(doc)
        
        # First, remove ALL extra pricing tables from template (keep only Tables 0, 1)
        # Template has tables at indices 2 and 3 which are pricing tables we need to remove extras
        # We'll populate them fresh for each location
        
        # Find all 4-column pricing tables (indices 2+)
        pricing_table_indices = []
        for i, table in enumerate(tables):
            if i >= 2 and len(table.columns) == 4 and len(table.rows) >= 8:
                pricing_table_indices.append(i)
        
        if not pricing_table_indices:
            return
        
        # Keep only the first pricing table, remove the rest
        first_pricing_table = tables[pricing_table_indices[0]]
        
        # Remove extra pricing tables (in reverse order to avoid index shifting)
        for idx in reversed(pricing_table_indices[1:]):
            try:
                tbl_element = tables[idx]._element
                parent = tbl_element.getparent()
                if parent is not None:
                    parent.remove(tbl_element)
            except:
                pass
        
        # Also remove any "PRICING DETAILS – BHIWANDI" or similar headers
        paragraphs_to_clear = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.upper()
            if i != pricing_header_idx and 'PRICING' in text and 'DETAILS' in text:
                paragraphs_to_clear.append(i)
        
        for idx in reversed(paragraphs_to_clear):
            doc.paragraphs[idx].text = ""  # Clear the extra headers
        
        # Now populate: First location uses existing table
        first_location = locations[0]
        
        # Update the header with first location name
        if pricing_header_idx >= 0:
            doc.paragraphs[pricing_header_idx].text = f"PRICING DETAILS – {first_location.location_name.upper()}"
        
        # Populate first location's table
        self._populate_table_with_items(first_pricing_table, first_location)
        
        # Get the element after the first pricing table to insert before
        last_element = first_pricing_table._element
        parent = last_element.getparent()
        
        # Find position of first pricing table in parent
        table_index = list(parent).index(last_element)
        insert_position = table_index + 1  # Insert after first table
        
        # Additional locations: insert right after
        for location in locations[1:]:
            # Create heading paragraph element
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            import copy
            
            # Create a new paragraph for heading
            p = OxmlElement('w:p')
            
            # Copy style from original header if possible
            if pricing_header_idx >= 0:
                original_header = doc.paragraphs[pricing_header_idx]._element
                if original_header.pPr is not None:
                    p.append(copy.deepcopy(original_header.pPr))
            
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = f"PRICING DETAILS – {location.location_name.upper()}"
            r.append(t)
            p.append(r)
            
            # Insert paragraph
            parent.insert(insert_position, p)
            insert_position += 1
            
            # Clone the first table to preserve ALL formatting (borders, shading, etc.)
            new_table_element = copy.deepcopy(first_pricing_table._element)
            parent.insert(insert_position, new_table_element)
            insert_position += 1
            
            # Create Table object from the new element
            from docx.table import Table
            new_table = Table(new_table_element, parent)
            
            # Prepare the cloned table: clear existing data rows (keep header/footer style)
            # We assume Row 0 is header. We need to reset content for items.
            # The template has specific rows. We will clear items and re-populate.
            
            # For the clone, we need to be careful. The template might have X rows.
            # We should clear rows 1 to N-3 (assuming Subtotal, GST, Grand Total are last 3)
            # Or simpler: Re-use _populate_table_with_items logic which handles row filling
            
            # However, the clone already HAS data from the first location (since we modified it first)
            # OR it has template data if we cloned before modification.
            # BUT we modified first_pricing_table ABOVE. So we are cloning the populated table.
            
            # Strategy: Clear the ITEM data rows, but keep the structure.
            # _populate_table_with_items overwrites data, but expects sufficient rows?
            # Or does it add rows?
            
            # Let's fix _populate_table_with_items to handle adding/removing rows to match item count
            # while preserving style.
            
            self._populate_cloned_table(new_table, location)

    def _populate_cloned_table(self, table, location):
        """Populate a cloned table with location data, adjusting rows as needed"""
        items = list(location.items.all().order_by('order'))
        
        # We assume the table has Header + Items + Subtotal + GST + Total
        # We want to identify the non-header, non-footer rows.
        # But simply: Row 0 is header. Last 3 rows are totals.
        # Everything in between is item slots.
        
        # Calculate needed item rows vs available item rows
        # Initial rows in the cloned table (which is a copy of populated Table 1)
        # Table 1 now has: 1 Header + N items (from loc 1) + 3 Totals
        
        current_rows = len(table.rows)
        footer_rows_count = 3 # Subtotal, GST, Total
        header_rows_count = 1
        
        available_item_rows = current_rows - header_rows_count - footer_rows_count
        needed_item_rows = len(items)
        
        # Adjust rows
        if needed_item_rows > available_item_rows:
            # We need more rows. Clone one of the item rows to preserve style?
            # Or just add rows. Adding rows usually lacks style.
            # Best to clone the LAST item row (before footer)
            insertion_idx = current_rows - footer_rows_count
            reference_row = table.rows[header_rows_count] # Use first item row as template
            
            for _ in range(needed_item_rows - available_item_rows):
                # Copy row element
                new_row_element = copy.deepcopy(reference_row._element)
                # clear text
                for cell in new_row_element.xpath(".//w:t"):
                    cell.text = ""
                # Insert before the footer
                table._element.insert(insertion_idx, new_row_element)
                insertion_idx += 1
                
        elif needed_item_rows < available_item_rows:
             # Remove excess rows
             rows_to_remove = available_item_rows - needed_item_rows
             # Remove from the end of the item section
             start_remove_idx = header_rows_count + needed_item_rows
             for _ in range(rows_to_remove):
                 row_elm = table.rows[start_remove_idx]._element
                 table._element.remove(row_elm)

        # Now we have exact number of rows. Populate them.
        for i, item in enumerate(items):
            row = table.rows[header_rows_count + i]
            self._fill_item_row(row, item)
            
        # Update totals
        self._update_totals_in_table(table, location)
    
    def _populate_table_with_items(self, table, location):
        """Populate an existing table with items from a location"""
        items = list(location.items.all().order_by('order'))
        
        # Skip header row (index 0), populate item rows
        for row_idx, item in enumerate(items):
            target_row_idx = row_idx + 1  # Skip header
            
            if target_row_idx < len(table.rows):
                row = table.rows[target_row_idx]
                self._fill_item_row(row, item)
        
        # Update totals rows
        self._update_totals_in_table(table, location)
    
    def _fill_item_row(self, row, item):
        """Fill a table row with item data"""
        cells = row.cells
        
        # Column 0: Description
        cells[0].text = item.display_description
        
        # Column 1: Unit Cost
        if str(item.unit_cost).lower() == 'at actual':
            cells[1].text = "At Actual"
        else:
            try:
                cost = Decimal(str(item.unit_cost))
                cells[1].text = f"₹ {cost:,.2f}"
            except:
                cells[1].text = str(item.unit_cost)
        
        # Column 2: Quantity
        if str(item.quantity).lower() == 'at actual':
            cells[2].text = "At Actual"
        else:
            cells[2].text = str(item.quantity)
        
        # Column 3: Total
        if item.is_calculated:
            cells[3].text = f"₹ {item.total:,.2f}"
        else:
            cells[3].text = "N/A"
    
    def _update_totals_in_table(self, table, location):
        """Update subtotal, GST, and grand total rows"""
        for row in table.rows:
            if len(row.cells) >= 4:
                first_cell_text = row.cells[0].text.lower().strip()
                
                # Check each cell for total keywords
                for cell in row.cells:
                    cell_text = cell.text.lower().strip()
                    if 'subtotal' in cell_text or 'sub-total' in cell_text:
                        # Find the last cell and set subtotal
                        row.cells[-1].text = f"₹ {location.subtotal:,.2f}"
                        break
                    elif 'gst' in cell_text:
                        row.cells[-1].text = f"₹ {location.gst_amount:,.2f}"
                        break
                    elif 'grand total' in cell_text:
                        row.cells[-1].text = f"₹ {location.grand_total:,.2f}"
                        break
    
    def _create_pricing_table_structure(self, table, location):
        """Create a pricing table structure for a location with proper styling"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Helper to set cell background
        def set_cell_background(cell, color_hex):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)

        items = list(location.items.all().order_by('order'))
        
        # Set header row
        header_row = table.rows[0]
        header_cells = header_row.cells
        header_cells[0].text = "Item Description"
        header_cells[1].text = "Unit Cost (₹)"
        header_cells[2].text = "Quantity"
        header_cells[3].text = "Total (₹)"
        
        # Style header: Red background, White bold text
        for cell in header_cells:
            set_cell_background(cell, 'C00000')  # Dark Red
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White
        
        # Add item rows
        for item in items:
            row = table.add_row()
            self._fill_item_row(row, item)
        
        # Add subtotal row
        subtotal_row = table.add_row()
        # Merge first 3 cells for label
        subtotal_label = subtotal_row.cells[0].merge(subtotal_row.cells[1]).merge(subtotal_row.cells[2])
        subtotal_label.text = "Subtotal"
        subtotal_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtotal_label.paragraphs[0].runs[0].font.bold = True
        
        subtotal_row.cells[3].text = f"₹ {location.subtotal:,.2f}"
        subtotal_row.cells[3].paragraphs[0].runs[0].font.bold = True
        
        # Add GST row
        gst_row = table.add_row()
        gst_label = gst_row.cells[0].merge(gst_row.cells[1]).merge(gst_row.cells[2])
        gst_label.text = "GST @ 18%"
        gst_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        gst_row.cells[3].text = f"₹ {location.gst_amount:,.2f}"
        
        # Add grand total row
        total_row = table.add_row()
        total_label = total_row.cells[0].merge(total_row.cells[1]).merge(total_row.cells[2])
        total_label.text = "GRAND TOTAL"
        total_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_row.cells[3].text = f"₹ {location.grand_total:,.2f}"
        
        # Style Grand Total: Red background, White bold text
        for cell in total_row.cells:
            set_cell_background(cell, 'C00000')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Try to apply grid style for borders
        try:
            table.style = 'Table Grid'
        except:
            pass


def generate_quotation_docx(quotation):
    """
    Helper function to generate DOCX quotation
    
    Args:
        quotation: Quotation model instance
        
    Returns:
        str: Path to generated DOCX file
    """
    generator = QuotationDocxGenerator(quotation)
    return generator.generate()
