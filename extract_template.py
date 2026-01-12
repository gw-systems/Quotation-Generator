from docx import Document

doc = Document("Godamwale-Quotation-Template.docx")

print("=== PARAGRAPHS ===")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        print(" | ".join(row_data))
